"""
导出服务 - Word/PDF 生成（V1.1优化版）

优化内容：
1. 标准化文档格式配置
2. 优化标题层级和间距
3. 增加人物外观描述导出
4. 规范化剧本格式
"""
import os
import zipfile
from datetime import datetime
from typing import List, Dict, Any, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from loguru import logger

from app.services.project import ProjectService


class ExportFormatConfig:
    """
    导出格式配置（V1.1新增）

    定义文档中各种元素的格式规范
    """

    # 字体设置
    TITLE_FONT = "黑体"
    SUBTITLE_FONT = "楷体"
    BODY_FONT = "宋体"
    SCENE_FONT = "黑体"

    # 字号设置
    TITLE_SIZE = Pt(22)
    SUBTITLE_SIZE = Pt(16)
    HEADING1_SIZE = Pt(18)
    HEADING2_SIZE = Pt(14)
    HEADING3_SIZE = Pt(12)
    BODY_SIZE = Pt(12)
    SCENE_SIZE = Pt(11)

    # 颜色设置
    TITLE_COLOR = RGBColor(0x1a, 0x1a, 0x1a)  # 深灰色
    SUBTITLE_COLOR = RGBColor(0x33, 0x33, 0x33)
    HOOK_COLOR = RGBColor(0xc0, 0x39, 0x2b)  # 红色

    # 行距设置
    LINE_SPACING = 1.5
    PARAGRAPH_SPACE_BEFORE = Pt(6)
    PARAGRAPH_SPACE_AFTER = Pt(6)

    # 页边距（厘米）
    PAGE_MARGIN_TOP = 2.54
    PAGE_MARGIN_BOTTOM = 2.54
    PAGE_MARGIN_LEFT = 3.17
    PAGE_MARGIN_RIGHT = 3.17


class ExportService:
    """导出服务（V1.1优化版）"""

    @staticmethod
    def _setup_document_styles(doc: Document) -> None:
        """
        设置文档样式（V1.1新增）

        Args:
            doc: Document 对象
        """
        # 设置页边距
        for section in doc.sections:
            section.top_margin = Cm(ExportFormatConfig.PAGE_MARGIN_TOP)
            section.bottom_margin = Cm(ExportFormatConfig.PAGE_MARGIN_BOTTOM)
            section.left_margin = Cm(ExportFormatConfig.PAGE_MARGIN_LEFT)
            section.right_margin = Cm(ExportFormatConfig.PAGE_MARGIN_RIGHT)

        # 设置 Normal 样式
        style = doc.styles['Normal']
        style.font.name = ExportFormatConfig.BODY_FONT
        style.font.size = ExportFormatConfig.BODY_SIZE
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    @staticmethod
    def _set_run_font(run, font_name: str, font_size: Pt, bold: bool = False) -> None:
        """
        设置文本格式

        Args:
            run: Run 对象
            font_name: 字体名称
            font_size: 字号
            bold: 是否加粗
        """
        run.font.name = font_name
        run.font.size = font_size
        run.bold = bold
        # 设置中文字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    @staticmethod
    def _add_formatted_paragraph(
        doc: Document,
        text: str,
        font_name: str,
        font_size: Pt,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        bold: bool = False,
        space_before: Pt = None,
        space_after: Pt = None
    ) -> None:
        """
        添加格式化段落

        Args:
            doc: Document 对象
            text: 文本内容
            font_name: 字体名称
            font_size: 字号
            alignment: 对齐方式
            bold: 是否加粗
            space_before: 段前间距
            space_after: 段后间距
        """
        p = doc.add_paragraph()
        run = p.add_run(text)
        ExportService._set_run_font(run, font_name, font_size, bold)

        p.alignment = alignment

        if space_before:
            p.paragraph_format.space_before = space_before
        if space_after:
            p.paragraph_format.space_after = space_after

        return p

    @staticmethod
    def export_to_docx(
        project_id: str,
        contents: List[str],
        episodes: str = "all"
    ) -> str:
        """
        导出为 Word 文档（V1.1优化版）

        Args:
            project_id: 项目ID
            contents: 导出内容类型列表 (synopsis, characters, outlines, scripts)
            episodes: 集数范围 (all/1-30/[1,2,3])

        Returns:
            文件路径
        """
        logger.info(f"导出 Word: {project_id}, contents={contents}, episodes={episodes}")

        project = ProjectService.get_project(project_id)
        if not project:
            raise ValueError("项目不存在")

        # 创建文档
        doc = Document()

        # 设置文档样式
        ExportService._setup_document_styles(doc)

        # ═══════════════════════════════════════════════════════════
        # 封面
        # ═══════════════════════════════════════════════════════════

        # 添加空白行
        for _ in range(3):
            doc.add_paragraph()

        # 主标题
        title_text = project.get("story_title") or project.get("name", "未命名剧本")
        ExportService._add_formatted_paragraph(
            doc, title_text,
            ExportFormatConfig.TITLE_FONT,
            ExportFormatConfig.TITLE_SIZE,
            WD_ALIGN_PARAGRAPH.CENTER,
            bold=True
        )

        # 副标题（一句话简介）
        if project.get("one_liner"):
            doc.add_paragraph()
            ExportService._add_formatted_paragraph(
                doc, project["one_liner"],
                ExportFormatConfig.SUBTITLE_FONT,
                ExportFormatConfig.SUBTITLE_SIZE,
                WD_ALIGN_PARAGRAPH.CENTER
            )

        # 添加空白行
        for _ in range(6):
            doc.add_paragraph()

        # 导出信息
        total_episodes = len(project.get("episode_outlines", []))
        info_lines = [
            f"总集数：{total_episodes} 集",
            f"题材类型：{project.get('requirements', {}).get('genre', '待定')}",
            f"风格基调：{project.get('requirements', {}).get('style', '待定')}",
            f"",
            f"导出时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}",
            f"生成工具：剧作大师 ScriptMaster V1.1"
        ]

        for line in info_lines:
            ExportService._add_formatted_paragraph(
                doc, line,
                ExportFormatConfig.BODY_FONT,
                Pt(11),
                WD_ALIGN_PARAGRAPH.CENTER
            )

        doc.add_page_break()

        # ═══════════════════════════════════════════════════════════
        # 目录页
        # ═══════════════════════════════════════════════════════════

        ExportService._add_formatted_paragraph(
            doc, "目 录",
            ExportFormatConfig.TITLE_FONT,
            Pt(18),
            WD_ALIGN_PARAGRAPH.CENTER,
            bold=True
        )
        doc.add_paragraph()

        toc_items = []
        if "synopsis" in contents and project.get("story_synopsis"):
            toc_items.append("一、故事梗概")
        if "characters" in contents and project.get("character_profiles"):
            toc_items.append("二、人物小传")
        if "outlines" in contents and project.get("episode_outlines"):
            toc_items.append("三、分集大纲")
        if "scripts" in contents and project.get("scripts"):
            toc_items.append("四、剧本正文")

        for item in toc_items:
            ExportService._add_formatted_paragraph(
                doc, item,
                ExportFormatConfig.BODY_FONT,
                ExportFormatConfig.BODY_SIZE
            )

        doc.add_page_break()

        # ═══════════════════════════════════════════════════════════
        # 一、故事梗概
        # ═══════════════════════════════════════════════════════════

        if "synopsis" in contents and project.get("story_synopsis"):
            ExportService._add_formatted_paragraph(
                doc, "一、故事梗概",
                ExportFormatConfig.TITLE_FONT,
                ExportFormatConfig.HEADING1_SIZE,
                bold=True
            )
            doc.add_paragraph()

            # 详细梗概
            if project.get("story_synopsis"):
                for para in project["story_synopsis"].split('\n'):
                    if para.strip():
                        ExportService._add_formatted_paragraph(
                            doc, para.strip(),
                            ExportFormatConfig.BODY_FONT,
                            ExportFormatConfig.BODY_SIZE,
                            space_after=Pt(12)
                        )

            # 核心卖点
            if project.get("selling_points"):
                doc.add_paragraph()
                ExportService._add_formatted_paragraph(
                    doc, "【核心卖点】",
                    ExportFormatConfig.SUBTITLE_FONT,
                    ExportFormatConfig.HEADING3_SIZE,
                    bold=True
                )
                for point in project["selling_points"]:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f" {point}")

            doc.add_page_break()

        # ═══════════════════════════════════════════════════════════
        # 二、人物小传
        # ═══════════════════════════════════════════════════════════

        if "characters" in contents and project.get("character_profiles"):
            ExportService._add_formatted_paragraph(
                doc, "二、人物小传",
                ExportFormatConfig.TITLE_FONT,
                ExportFormatConfig.HEADING1_SIZE,
                bold=True
            )
            doc.add_paragraph()

            for i, char in enumerate(project["character_profiles"], 1):
                char_name = ExportService._safe_get_field(char, 'name', 'name', '未知')
                char_role = ExportService._safe_get_field(char, 'role', 'role', '')

                # 人物标题
                role_text = f"（{char_role}）" if char_role else ""
                ExportService._add_formatted_paragraph(
                    doc, f"{i}. {char_name}{role_text}",
                    ExportFormatConfig.SUBTITLE_FONT,
                    ExportFormatConfig.HEADING2_SIZE,
                    bold=True
                )

                # 外观描述（V1.1新增）
                appearance = char.get('appearance', {})
                if appearance:
                    ExportService._add_formatted_paragraph(
                        doc, "【外观形象】",
                        ExportFormatConfig.BODY_FONT,
                        Pt(11),
                        bold=True,
                        space_before=Pt(6)
                    )
                    # 更安全的处理方式
                    appearance_parts = []
                    if appearance.get('height'):
                        appearance_parts.append(f"身高{appearance['height']}")
                    if appearance.get('build'):
                        appearance_parts.append(appearance['build'])
                    if appearance.get('hair'):
                        appearance_parts.append(appearance['hair'])
                    if appearance.get('clothing_style'):
                        appearance_parts.append(f"穿着{appearance['clothing_style']}")
                    if appearance.get('distinctive_features'):
                        appearance_parts.append(appearance['distinctive_features'])

                    if appearance_parts:
                        ExportService._add_formatted_paragraph(
                            doc, "，".join(appearance_parts),
                            ExportFormatConfig.BODY_FONT,
                            ExportFormatConfig.BODY_SIZE
                        )

                # 基本信息表格
                info_table = doc.add_table(rows=4, cols=2)
                info_table.style = 'Light Grid Accent 1'
                info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                info_data = [
                    ("年龄", ExportService._safe_get_field(char, 'age', 'age', '')),
                    ("性格", ExportService._safe_get_field(char, 'personality', 'personality', '')),
                    ("背景", ExportService._safe_get_field(char, 'background', 'background', '')),
                    ("目标", ExportService._safe_get_field(char, 'goal', 'goal', '')),
                ]

                for row_idx, (key, value) in enumerate(info_data):
                    cell0 = info_table.rows[row_idx].cells[0]
                    cell1 = info_table.rows[row_idx].cells[1]
                    cell0.text = key
                    cell1.text = str(value) if value else ""

                # 记忆点
                memory_point = ExportService._safe_get_field(char, 'memoryPoint', 'memory_point', '')
                if memory_point:
                    doc.add_paragraph()
                    ExportService._add_formatted_paragraph(
                        doc, f"记忆点：{memory_point}",
                        ExportFormatConfig.BODY_FONT,
                        ExportFormatConfig.BODY_SIZE
                    )

                doc.add_paragraph()  # 空行

            doc.add_page_break()

        # ═══════════════════════════════════════════════════════════
        # 三、分集大纲
        # ═══════════════════════════════════════════════════════════

        if "outlines" in contents and project.get("episode_outlines"):
            ExportService._add_formatted_paragraph(
                doc, "三、分集大纲",
                ExportFormatConfig.TITLE_FONT,
                ExportFormatConfig.HEADING1_SIZE,
                bold=True
            )
            doc.add_paragraph()

            episode_list = ExportService._parse_episodes(episodes, len(project["episode_outlines"]))

            for ep_num in episode_list:
                outline = next(
                    (o for o in project["episode_outlines"]
                     if ExportService._safe_get_field(o, 'episodeNumber', 'episode_number') == ep_num),
                    None
                )
                if outline:
                    is_checkpoint = ExportService._safe_get_field(outline, 'isCheckpoint', 'is_checkpoint', False)
                    summary = ExportService._safe_get_field(outline, 'summary', 'summary', '')
                    hook = ExportService._safe_get_field(outline, 'hook', 'hook', '')

                    # 集标题
                    title = f"第{ep_num}集"
                    if is_checkpoint:
                        title += " ★ 付费卡点"

                    ExportService._add_formatted_paragraph(
                        doc, title,
                        ExportFormatConfig.SUBTITLE_FONT,
                        ExportFormatConfig.HEADING3_SIZE,
                        bold=True
                    )

                    # 剧情概要
                    if summary:
                        p = doc.add_paragraph()
                        p.add_run("剧情概要：").bold = True
                        p.add_run(summary)

                    # 卡点/钩子
                    if hook:
                        p = doc.add_paragraph()
                        p.add_run("卡点/钩子：").bold = True
                        run = p.add_run(hook)
                        run.font.color.rgb = ExportFormatConfig.HOOK_COLOR

                    doc.add_paragraph()  # 空行

            doc.add_page_break()

        # ═══════════════════════════════════════════════════════════
        # 四、剧本正文
        # ═══════════════════════════════════════════════════════════

        if "scripts" in contents and project.get("scripts"):
            ExportService._add_formatted_paragraph(
                doc, "四、剧本正文",
                ExportFormatConfig.TITLE_FONT,
                ExportFormatConfig.HEADING1_SIZE,
                bold=True
            )
            doc.add_paragraph()

            episode_list = ExportService._parse_episodes(episodes, len(project["scripts"]))

            for ep_num in episode_list:
                script = next(
                    (s for s in project["scripts"]
                     if ExportService._safe_get_field(s, 'episodeNumber', 'episode_number') == ep_num),
                    None
                )
                if script:
                    # 集标题
                    title = f"第{ep_num}集"
                    script_title = ExportService._safe_get_field(script, 'title', 'title', '')
                    if script_title:
                        title += f" {script_title}"

                    ExportService._add_formatted_paragraph(
                        doc, title,
                        ExportFormatConfig.SUBTITLE_FONT,
                        ExportFormatConfig.HEADING2_SIZE,
                        WD_ALIGN_PARAGRAPH.CENTER,
                        bold=True
                    )

                    # 字数信息
                    word_count = ExportService._safe_get_field(script, 'wordCount', 'word_count', 0)
                    if word_count:
                        ExportService._add_formatted_paragraph(
                            doc, f"（字数：{word_count}字）",
                            ExportFormatConfig.BODY_FONT,
                            Pt(10),
                            WD_ALIGN_PARAGRAPH.CENTER
                        )

                    doc.add_paragraph()

                    # 正文
                    content = ExportService._safe_get_field(script, 'content', 'content', '')
                    if content:
                        lines = content.split('\n')
                        for line in lines:
                            if line.strip():
                                # 场景标题 - 居中
                                if line.startswith("场景："):
                                    ExportService._add_formatted_paragraph(
                                        doc, line.strip(),
                                        ExportFormatConfig.SCENE_FONT,
                                        ExportFormatConfig.SCENE_SIZE,
                                        WD_ALIGN_PARAGRAPH.CENTER,
                                        space_before=Pt(12)
                                    )
                                # 人物名 - 居中加粗
                                elif line.startswith("**") and ":**" in line:
                                    # 格式: **人物名（情绪）：**
                                    clean_line = line.replace("**", "")
                                    ExportService._add_formatted_paragraph(
                                        doc, clean_line,
                                        ExportFormatConfig.BODY_FONT,
                                        ExportFormatConfig.BODY_SIZE,
                                        WD_ALIGN_PARAGRAPH.CENTER,
                                        bold=True,
                                        space_before=Pt(6)
                                    )
                                # 动作描述 - 斜体
                                elif line.startswith("▶"):
                                    p = doc.add_paragraph()
                                    run = p.add_run(line[1:].strip())
                                    run.italic = True
                                    run.font.size = Pt(11)
                                # 卡点标记 - 红色
                                elif line.startswith("【本集卡点"):
                                    p = doc.add_paragraph()
                                    run = p.add_run(line.strip())
                                    run.font.color.rgb = ExportFormatConfig.HOOK_COLOR
                                    run.font.size = Pt(11)
                                # 普通台词
                                else:
                                    ExportService._add_formatted_paragraph(
                                        doc, line,
                                        ExportFormatConfig.BODY_FONT,
                                        ExportFormatConfig.BODY_SIZE
                                    )
                            else:
                                doc.add_paragraph()  # 空行

                    doc.add_paragraph()
                    # 分隔线
                    p = doc.add_paragraph()
                    p.add_run("━" * 40)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

        # 保存文件
        output_dir = os.path.join(os.getcwd(), "exports")
        os.makedirs(output_dir, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        project_name = project.get("story_title") or project.get("name", "未命名")
        # 清理文件名中的非法字符
        safe_name = "".join(c for c in project_name if c.isalnum() or c in (' ', '-', '_')).strip()
        filename = f"{safe_name}_{timestamp}.docx"
        filepath = os.path.join(output_dir, filename)

        doc.save(filepath)
        logger.info(f"Word文档已保存: {filepath}")

        return filepath

    @staticmethod
    def export_to_pdf(
        project_id: str,
        contents: List[str],
        episodes: str = "all"
    ) -> str:
        """
        导出为 PDF

        Args:
            project_id: 项目ID
            contents: 导出内容类型列表
            episodes: 集数范围

        Returns:
            文件路径
        """
        # 暂时通过 Word 转换
        # TODO: 使用 weasyprint 实现原生 PDF 导出
        logger.info(f"导出 PDF: {project_id}, contents={contents}, episodes={episodes}")

        # 先生成 Word
        docx_path = ExportService.export_to_docx(project_id, contents, episodes)

        # 提示用户 Word 转 PDF 的方法
        pdf_path = docx_path.replace(".docx", ".pdf")

        return pdf_path

    @staticmethod
    def export_to_zip(
        project_id: str,
        contents: List[str],
        episodes: str = "all"
    ) -> str:
        """
        导出为 ZIP 压缩包

        Args:
            project_id: 项目ID
            contents: 导出内容类型列表
            episodes: 集数范围

        Returns:
            文件路径
        """
        logger.info(f"导出 ZIP: {project_id}, contents={contents}, episodes={episodes}")

        project = ProjectService.get_project(project_id)
        if not project:
            raise ValueError("项目不存在")

        # 创建临时目录
        output_dir = os.path.join(os.getcwd(), "exports")
        os.makedirs(output_dir, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        project_name = project.get("story_title") or project.get("name", "未命名")
        safe_name = "".join(c for c in project_name if c.isalnum() or c in (' ', '-', '_')).strip()
        zip_filename = f"{safe_name}_{timestamp}.zip"
        zip_filepath = os.path.join(output_dir, zip_filename)

        with zipfile.ZipFile(zip_filepath, 'w', zipfile.ZIP_DEFLATED) as zf:
            # 添加 Word 文档
            docx_path = ExportService.export_to_docx(project_id, contents, episodes)
            zf.write(docx_path, os.path.basename(docx_path))

        logger.info(f"ZIP包已保存: {zip_filepath}")
        return zip_filepath

    @staticmethod
    def _parse_episodes(episodes: str, total: int) -> List[int]:
        """
        解析集数范围

        Args:
            episodes: 集数范围字符串 (all/1-30/1,2,3)
            total: 总集数

        Returns:
            集数列表
        """
        if episodes == "all":
            return list(range(1, total + 1))

        if "-" in episodes:
            start, end = episodes.split("-")
            return list(range(int(start), int(end) + 1))

        if "," in episodes:
            return [int(x.strip()) for x in episodes.split(",")]

        try:
            return [int(episodes)]
        except ValueError:
            return list(range(1, total + 1))

    @staticmethod
    def _safe_get_field(obj: dict, camel_case: str, snake_case: str, default=None):
        """兼容获取字段（支持 camelCase 和 snake_case）"""
        if not obj:
            return default
        return obj.get(camel_case, obj.get(snake_case, default))
